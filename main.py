"""
Backend FastAPI pour le traitement de documents administratifs marocains
- Pipeline: PDF → Images → Prétraitement → OCR → NLP → Post-traitement
- Gestion d'erreurs complète
- Validation des fichiers uploadés
- Logs structurés pour le débogage
"""
import time
import uuid
import logging
import hashlib
import json
from io import BytesIO
from contextlib import asynccontextmanager
from pathlib import Path

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware

from utils.pdf_to_image import convert_pdf_to_images
from Services.preprocess import preprocess
from Services.ocr import run_ocr, get_reader, get_ocr_engine_name
from Services.nlp import extract_entities, get_ner
from Services.postprocess import clean_output

# Configuration des logs
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)

# Taille max du fichier (20 MB)
MAX_FILE_SIZE = 20 * 1024 * 1024
RESULTS_CACHE_DIR = Path(__file__).resolve().parent / "results_cache"
CACHE_VERSION = "business-schema-v7"
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def _cache_path(file_hash: str) -> Path:
    return RESULTS_CACHE_DIR / f"{file_hash}.json"


def _load_cached_result(file_hash: str):
    path = _cache_path(file_hash)
    if not path.exists():
        return None

    try:
        with path.open("r", encoding="utf-8") as f:
            cached = json.load(f)
        if cached.get("_cache_version") != CACHE_VERSION:
            return None
        cached.pop("_cache_version", None)
        cached.pop("_source_type", None)
        return cached
    except Exception as e:
        logger.warning(f"Cache illisible ({path.name}): {e}")
        return None


def _save_cached_result(file_hash: str, result: dict):
    RESULTS_CACHE_DIR.mkdir(exist_ok=True)
    path = _cache_path(file_hash)
    cache_payload = dict(result)
    cache_payload["_cache_version"] = CACHE_VERSION

    with path.open("w", encoding="utf-8") as f:
        json.dump(cache_payload, f, ensure_ascii=False, indent=2)


def _extract_docx_text(docx_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx est requis pour lire les fichiers DOCX") from e

    document = Document(BytesIO(docx_bytes))
    chunks = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip().replace("\n", " ")
                for cell in row.cells
                if cell.text.strip()
            ]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks)


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Pre-charger les modeles au demarrage."""
    logger.info("--- Chargement des modeles (2-3 min) ---")
    total = time.time()

    t = time.time()
    get_reader()
    logger.info(f"EasyOCR charge en {time.time() - t:.1f}s")

    t = time.time()
    get_ner()
    logger.info(f"BERT NER charge en {time.time() - t:.1f}s")

    logger.info(f"Serveur pret ! (total: {time.time() - total:.1f}s)")
    yield
    logger.info("Arret du serveur")


app = FastAPI(
    title="API Extraction Documents Marocains",
    description="OCR + NLP pour documents administratifs arabes/francais",
    version="2.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
async def health():
    return {"status": "ok", "version": "2.0.0"}


@app.get("/diag")
async def diag():
    """Diagnostic: vérifie que tous les modules sont bien chargés."""
    status = {"ocr_engine": get_ocr_engine_name()}
    
    # Test arabic_utils
    try:
        from Services.arabic_utils import normalize_for_matching, build_flexible_pattern
        status["arabic_utils"] = "OK"
        status["normalize_test"] = normalize_for_matching("تجربة")
        status["flex_test"] = build_flexible_pattern("وزارة")[:30] + "..."
    except Exception as e:
        status["arabic_utils"] = f"ERREUR: {e}"
    
    # Test rapidfuzz
    try:
        import rapidfuzz
        status["rapidfuzz"] = f"OK (version: {rapidfuzz.__version__})"
    except ImportError:
        status["rapidfuzz"] = "NON INSTALLE - pip install rapidfuzz"
    except Exception as e:
        status["rapidfuzz"] = f"ERREUR: {e}"
    
    # Test regex
    try:
        import regex
        status["regex"] = f"OK (version: {regex.__version__})"
    except ImportError:
        status["regex"] = "NON INSTALLE - pip install regex"
    except Exception as e:
        status["regex"] = f"ERREUR: {e}"
    
    # Test fuzzy_match
    try:
        from Services.fuzzy_match import ExtractionResult, fuzzy_find_keyword
        status["fuzzy_match"] = "OK"
    except Exception as e:
        status["fuzzy_match"] = f"ERREUR: {e}"
    
    # Test postprocess
    try:
        from Services.postprocess import clean_output
        status["postprocess"] = "OK"
    except Exception as e:
        status["postprocess"] = f"ERREUR: {e}"
    
    # Test nlp
    try:
        from Services.nlp import extract_entities
        status["nlp"] = "OK"
    except Exception as e:
        status["nlp"] = f"ERREUR: {e}"
    
    return status


@app.post("/extract")
async def extract(file: UploadFile = File(...)):
    """
    Endpoint principal: extrait les informations d'un PDF arabe.
    
    Retourne: montant, date, organisations, personnes, lieux, reference, raw_text
    """
    request_id = str(uuid.uuid4())[:8]
    logger.info(f"[{request_id}] === Nouvelle requete: {file.filename} ===")
    total_start = time.time()

    # --- Validation ---
    if not file.filename:
        raise HTTPException(status_code=400, detail="Nom de fichier manquant")

    file_extension = Path(file.filename).suffix.lower()
    if file_extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Format non supporte: {file.filename}. "
                "Formats acceptes: PDF, DOCX."
            )
        )

    try:
        pdf_bytes = await file.read()
    except Exception as e:
        logger.error(f"[{request_id}] Erreur lecture fichier: {e}")
        raise HTTPException(status_code=400, detail="Impossible de lire le fichier")

    if len(pdf_bytes) == 0:
        raise HTTPException(status_code=400, detail="Le fichier est vide")

    if len(pdf_bytes) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"Fichier trop volumineux ({len(pdf_bytes) // 1024 // 1024} MB). Max: {MAX_FILE_SIZE // 1024 // 1024} MB"
        )

    file_hash = hashlib.sha256(pdf_bytes).hexdigest()
    cached_result = _load_cached_result(file_hash)
    if cached_result:
        logger.info(f"[{request_id}] Resultat deja enregistre, retour depuis le cache")
        return cached_result

    logger.info(f"[{request_id}] Fichier valide: {len(pdf_bytes) // 1024} KB")

    try:
        if file_extension == ".pdf":
            t = time.time()
            images = convert_pdf_to_images(pdf_bytes)
            logger.info(f"[{request_id}] 1/5 PDF -> {len(images)} images en {time.time() - t:.1f}s")

            t = time.time()
            processed_images = [preprocess(img) for img in images]
            logger.info(f"[{request_id}] 2/5 Pretraitement en {time.time() - t:.1f}s")

            t = time.time()
            text = run_ocr(processed_images)
            logger.info(f"[{request_id}] 3/5 OCR en {time.time() - t:.1f}s - {len(text)} caracteres")
            source_type = "pdf_ocr"
        else:
            t = time.time()
            text = _extract_docx_text(pdf_bytes)
            logger.info(f"[{request_id}] 1/3 DOCX -> texte en {time.time() - t:.1f}s - {len(text)} caracteres")
            source_type = "docx_text"

        if not text.strip():
            logger.warning(f"[{request_id}] Aucun texte detecte!")
            result = {
                "warning": "Aucun texte detecte dans le document",
                "source_type": source_type,
            }
            _save_cached_result(file_hash, result)
            return result

        # 4. NLP (BERT arabe)
        t = time.time()
        entities = extract_entities(text)
        logger.info(f"[{request_id}] 4/5 NLP en {time.time() - t:.1f}s — {len(entities)} entites")

        # 5. Post-traitement
        t = time.time()
        result = clean_output(text, entities)
        logger.info(f"[{request_id}] 5/5 Post-traitement en {time.time() - t:.1f}s")

        # Ajouter les metadonnees
        processing_time = round(time.time() - total_start, 1)

        logger.info(
            f"[{request_id}] === Termine en {processing_time}s === "
            f"Champs: {list(result.keys())}"
        )

        cache_result = dict(result)
        cache_result["_source_type"] = source_type
        _save_cached_result(file_hash, cache_result)
        return result

    except ValueError as e:
        logger.error(f"[{request_id}] Erreur de valeur: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"[{request_id}] Erreur interne: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Erreur lors du traitement: {str(e)}"
        )
